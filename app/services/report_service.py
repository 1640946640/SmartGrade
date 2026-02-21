from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class ReportGenerator:
    def generate_word_report(self, report_data, output_path):
        """
        生成Word格式的批改报告
        """
        doc = Document()
        
        # 标题
        heading = doc.add_heading('智能试卷批改报告', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. 批改概览
        doc.add_heading('一、批改概览', level=1)
        
        summary_table = doc.add_table(rows=1, cols=3)
        summary_table.style = 'Table Grid'
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = '总分'
        hdr_cells[1].text = '正确率'
        hdr_cells[2].text = '正确题数'
        
        # 设置表头背景色 (需要xml操作，暂略，仅加粗)
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        row_cells = summary_table.add_row().cells
        
        # 获取数据
        final_results = report_data.get('final_results', {})
        total_score = final_results.get('total_score', 0)
        max_total_score = final_results.get('max_total_score', 0)
        accuracy = final_results.get('accuracy', 0)
        correct_count = final_results.get('correct_count', 0)
        total_count = final_results.get('total_count', 0)
        
        row_cells[0].text = f"{total_score} / {max_total_score}"
        row_cells[1].text = f"{accuracy * 100:.1f}%"
        row_cells[2].text = f"{correct_count} / {total_count}"
        
        doc.add_paragraph() # 空行
        
        # 2. 详细批改结果
        doc.add_heading('二、详细批改结果', level=1)
        
        details = final_results.get('details', {})
        
        # 尝试智能排序
        def sort_key(k):
            # 尝试将 "1", "2" 转为数字， "一、填空题-1" 转为 ("一、填空题", 1)
            try:
                k_str = str(k)
                if '-' in k_str:
                    parts = k_str.rsplit('-', 1) # 从右边分割，确保只分割出最后的数字题号
                    group_part = parts[0]
                    num_part = parts[1]
                    
                    # 尝试将题号转为数字
                    try:
                        num = int(num_part)
                    except:
                        num = float('inf') # 如果不是数字，排在最后
                        
                    # 定义大题的排序权重（中文数字）
                    group_weights = {
                        "一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
                        "填空题": 1, "选择题": 2, "判断题": 3, "简答题": 4, "计算题": 5, "编程题": 6, "综合题": 7
                    }
                    
                    group_weight = 999
                    for key, weight in group_weights.items():
                        if key in group_part:
                            group_weight = weight
                            break
                            
                    return (group_weight, group_part, num)
                
                # 如果没有横杠，直接尝试转数字
                return (0, "", int(k_str))
            except:
                return (float('inf'), str(k), 0)
        
        sorted_keys = sorted(details.keys(), key=sort_key)
        
        for q_id in sorted_keys:
            item = details[q_id]
            
            # 题号标题
            group_name = item.get('group_name', '')
            sub_id = item.get('sub_id', '')
            
            if group_name and sub_id:
                q_title_text = f"{group_name} 第{sub_id}题"
            else:
                q_title_text = f"第 {q_id} 题"
                
            # 使用表格来布局题目信息
            q_table = doc.add_table(rows=1, cols=2)
            q_table.autofit = True
            
            # 左侧：题号
            cell_left = q_table.rows[0].cells[0]
            p = cell_left.paragraphs[0]
            run = p.add_run(q_title_text)
            run.bold = True
            run.font.size = Pt(12)
            
            # 右侧：得分和状态
            cell_right = q_table.rows[0].cells[1]
            p = cell_right.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            score = item.get('score', 0)
            max_score = item.get('max_score', 0)
            is_correct = item.get('is_correct', False)
            
            score_run = p.add_run(f"得分: {score}/{max_score}  ")
            status_run = p.add_run("【正确】" if is_correct else "【错误】")
            status_run.font.color.rgb = RGBColor(0, 128, 0) if is_correct else RGBColor(255, 0, 0)
            status_run.bold = True
            
            # 详细分析内容
            analysis = item.get('analysis', '')
            comment = item.get('comment', '')
            
            # 如果没有analysis，使用comment
            content = analysis if analysis else comment
            if not content:
                content = "暂无详细分析"
                
            doc.add_paragraph("答题分析:", style='List Bullet')
            analysis_p = doc.add_paragraph(content)
            analysis_p.paragraph_format.left_indent = Inches(0.25)
            
            # 添加分割线
            doc.add_paragraph("_" * 40)
            
        # 保存
        doc.save(output_path)
        return output_path
